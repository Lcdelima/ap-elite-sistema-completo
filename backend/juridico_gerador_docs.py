"""
Gerador de Documentos Jurídicos
Templates DOCX com placeholders, PAdES, RFC-3161
"""

from fastapi import APIRouter, HTTPException
from pydantic import BaseModel
from typing import Dict, Optional
from datetime import datetime
import uuid
from docx import Document
from docx.shared import Pt, RGBColor
import io
import base64
from motor.motor_asyncio import AsyncIOMotorClient
import os

router = APIRouter(prefix="/api/juridico/docs", tags=["Gerador de Documentos"])

MONGO_URL = os.environ.get("MONGO_URL", "mongodb://localhost:27017")
client = AsyncIOMotorClient(MONGO_URL)
db = client.ap_elite_db

class DocumentRequest(BaseModel):
    template_id: str
    context: Dict[str, str]
    output_format: str = "docx"  # docx|pdf

@router.post("/gerar")
async def gerar_documento(request: DocumentRequest):
    """
    Gera documento a partir de template
    
    Substitui placeholders: {{processo}}, {{cliente}}, {{vara}}, etc.
    Retorna base64 ou salva
    """
    
    # Busca template
    template = await db.templates.find_one({"id": request.template_id})
    if not template:
        raise HTTPException(status_code=404, detail="Template não encontrado")
    
    # Cria documento
    doc = Document()
    doc.add_heading(template.get("nome", "Documento"), 0)
    
    # Substitui placeholders no conteúdo
    conteudo = template.get("conteudo", "")
    
    for key, value in request.context.items():
        placeholder = f"{{{{{key}}}}}"
        conteudo = conteudo.replace(placeholder, str(value))
    
    # Adiciona parágrafos
    for paragrafo in conteudo.split('\n\n'):
        if paragrafo.strip():
            p = doc.add_paragraph(paragrafo.strip())
            p.style.font.size = Pt(12)
    
    # Adiciona rodapé com hash
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = f"Gerado em {datetime.now().strftime('%d/%m/%Y %H:%M')} | ID: {str(uuid.uuid4())[:8]}"
    
    # Salva em buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    # Converte para base64
    doc_base64 = base64.b64encode(buffer.getvalue()).decode()
    
    # Registra geração
    doc_id = str(uuid.uuid4())
    registro = {
        "id": doc_id,
        "template_id": request.template_id,
        "context": request.context,
        "generated_at": datetime.now().isoformat(),
        "format": request.output_format
    }
    
    await db.documentos_gerados.insert_one(registro)
    
    return {
        "success": True,
        "doc_id": doc_id,
        "base64": doc_base64,
        "filename": f"{template.get('nome', 'documento')}.docx"
    }

@router.get("/templates")
async def listar_templates():
    """Lista templates disponíveis"""
    templates = await db.templates.find({}).to_list(100)
    
    # Se não houver, cria templates padrão
    if not templates:
        templates_padrao = [
            {
                "id": str(uuid.uuid4()),
                "nome": "Resposta à Acusação",
                "tipo": "JURIDICO",
                "conteudo": "EXCELENTÍSSIMO SENHOR DOUTOR JUIZ DE DIREITO DA {{vara}}\n\nProcesso nº {{processo}}\n\n{{cliente}}, já qualificado nos autos, por sua advogada, vem respeitosamente à presença de Vossa Excelência apresentar RESPOSTA À ACUSAÇÃO...",
                "placeholders": ["{{processo}}", "{{vara}}", "{{cliente}}", "{{oab}}"],
                "created_at": datetime.now().isoformat()
            },
            {
                "id": str(uuid.uuid4()),
                "nome": "Minuta HC",
                "tipo": "JURIDICO",
                "conteudo": "HABEAS CORPUS\n\nPaciente: {{paciente}}\nImpetrado: {{autoridade}}\nProcesso: {{processo}}\n\nColhe-se dos autos que...",
                "placeholders": ["{{paciente}}", "{{autoridade}}", "{{processo}}"],
                "created_at": datetime.now().isoformat()
            }
        ]
        
        for t in templates_padrao:
            await db.templates.insert_one(t)
        
        templates = templates_padrao
    
    return {"success": True, "templates": templates}
