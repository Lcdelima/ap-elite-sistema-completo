"""
Sistema de Geração de Documentos Jurídicos
Gera documentos usando templates e IA
"""

from fastapi import APIRouter, HTTPException
from pydantic import BaseModel
from typing import List, Optional, Dict, Any
import os
from datetime import datetime
from motor.motor_asyncio import AsyncIOMotorClient
from ai_orchestrator import ai_orchestrator
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

router = APIRouter(prefix="/api/templates", tags=["Document Templates"])

# MongoDB connection
MONGO_URL = os.environ.get("MONGO_URL")
client = AsyncIOMotorClient(MONGO_URL)
db = client[os.environ.get("DB_NAME", "test_database")]

TEMPLATES_DIR = "/app/backend/templates"
GENERATED_DIR = "/app/backend/generated_documents"
os.makedirs(TEMPLATES_DIR, exist_ok=True)
os.makedirs(GENERATED_DIR, exist_ok=True)

class TemplateInfo(BaseModel):
    id: str
    name: str
    description: str
    category: str
    fields: List[str]
    structure: Dict[str, Any]

class GenerateDocumentRequest(BaseModel):
    template_id: str
    case_data: Dict[str, Any]
    use_ai_completion: bool = True
    ai_provider: Optional[str] = 'anthropic'  # Melhor para documentos

# Templates pré-definidos baseados nos documentos fornecidos
DOCUMENT_TEMPLATES = {
    'aij_roteiro': {
        'id': 'aij_roteiro',
        'name': 'Roteiro de Audiência de Instrução e Julgamento (AIJ)',
        'description': 'Template completo para condução de AIJ com fundamentos jurídicos',
        'category': 'audiencias',
        'fields': [
            'autos_numero',
            'processo',
            'data_audiencia',
            'juizo',
            'comarca',
            'defesa_tecnica',
            'assistente_tecnica',
            'ministerio_publico',
            'defensores',
            'reus',
            'vitimas',
            'testemunhas',
            'peritos'
        ],
        'structure': {
            'secoes': [
                'I. IDENTIFICAÇÃO DO PROCESSO',
                'II. IDENTIFICAÇÃO DAS PARTES',
                'III. LINHA DE INQUIRIÇÃO / AUDIÊNCIA',
                'IV. ANOTAÇÕES DE AUDIÊNCIA',
                'V. FUNDAMENTOS JURÍDICOS RELEVANTES',
                'VI. TRIAGEM E REFERÊNCIAS DOCUMENTAIS',
                'VII. CONCLUSÃO E PROVIDÊNCIAS PÓS-AUDIÊNCIA'
            ]
        }
    },
    'procuracao': {
        'id': 'procuracao',
        'name': 'Procuração Modelo Técnico Elite',
        'description': 'Procuração para representação em investigações técnicas',
        'category': 'procuracoes',
        'fields': [
            'outorgante_nome',
            'outorgante_cpf',
            'outorgante_rg',
            'outorgante_endereco',
            'outorgado_nome',
            'outorgado_oab',
            'poderes',
            'finalidade'
        ],
        'structure': {
            'secoes': [
                'OUTORGANTE',
                'OUTORGADO',
                'PODERES',
                'CLÁUSULAS ESPECIAIS',
                'FORO'
            ]
        }
    },
    'termo_confidencialidade': {
        'id': 'termo_confidencialidade',
        'name': 'Termo de Confidencialidade',
        'description': 'Termo de confidencialidade para casos sensíveis',
        'category': 'termos',
        'fields': [
            'contratante',
            'contratada',
            'objeto',
            'vigencia',
            'penalidades'
        ],
        'structure': {
            'secoes': [
                'PARTES',
                'OBJETO',
                'OBRIGAÇÕES DE CONFIDENCIALIDADE',
                'VIGÊNCIA',
                'PENALIDADES',
                'DISPOSIÇÕES FINAIS'
            ]
        }
    },
    'ata_reuniao': {
        'id': 'ata_reuniao',
        'name': 'Ata de Reunião',
        'description': 'Registro formal de reuniões de caso',
        'category': 'atas',
        'fields': [
            'data',
            'hora_inicio',
            'hora_fim',
            'local',
            'participantes',
            'pauta',
            'deliberacoes'
        ],
        'structure': {
            'secoes': [
                'IDENTIFICAÇÃO',
                'PARTICIPANTES',
                'PAUTA',
                'DESENVOLVIMENTO',
                'DELIBERAÇÕES',
                'ENCERRAMENTO'
            ]
        }
    },
    'relatorio_investigacao': {
        'id': 'relatorio_investigacao',
        'name': 'Relatório de Investigação',
        'description': 'Relatório técnico de investigação criminal',
        'category': 'relatorios',
        'fields': [
            'caso_numero',
            'data_inicio',
            'investigador',
            'objeto',
            'metodologia',
            'achados',
            'conclusoes',
            'recomendacoes'
        ],
        'structure': {
            'secoes': [
                '1. IDENTIFICAÇÃO',
                '2. OBJETO DA INVESTIGAÇÃO',
                '3. METODOLOGIA',
                '4. ANÁLISE E ACHADOS',
                '5. CONCLUSÕES',
                '6. RECOMENDAÇÕES',
                '7. ANEXOS'
            ]
        }
    },
    'analise_provas': {
        'id': 'analise_provas',
        'name': 'Roteiro de Análise de Provas',
        'description': 'Análise sistemática de provas e evidências',
        'category': 'pericia',
        'fields': [
            'caso_numero',
            'tipo_prova',
            'data_coleta',
            'cadeia_custodia',
            'metodologia_analise',
            'resultados',
            'conclusao_pericial'
        ],
        'structure': {
            'secoes': [
                'IDENTIFICAÇÃO DA PROVA',
                'CADEIA DE CUSTÓDIA',
                'METODOLOGIA DE ANÁLISE',
                'RESULTADOS',
                'CONCLUSÃO PERICIAL',
                'FUNDAMENTAÇÃO TÉCNICA'
            ]
        }
    }
}

@router.get("/list")
async def list_templates():
    """Lista todos os templates disponíveis"""
    return {
        "templates": list(DOCUMENT_TEMPLATES.values()),
        "total": len(DOCUMENT_TEMPLATES),
        "categories": list(set(t['category'] for t in DOCUMENT_TEMPLATES.values()))
    }

@router.get("/{template_id}")
async def get_template(template_id: str):
    """Obtém detalhes de um template"""
    if template_id not in DOCUMENT_TEMPLATES:
        raise HTTPException(status_code=404, detail="Template não encontrado")
    
    return DOCUMENT_TEMPLATES[template_id]

@router.post("/generate")
async def generate_document(request: GenerateDocumentRequest):
    """Gera documento baseado em template"""
    
    if request.template_id not in DOCUMENT_TEMPLATES:
        raise HTTPException(status_code=404, detail="Template não encontrado")
    
    template_info = DOCUMENT_TEMPLATES[request.template_id]
    
    # Se usar IA para completar campos faltantes
    if request.use_ai_completion:
        # Identificar campos faltantes
        missing_fields = set(template_info['fields']) - set(request.case_data.keys())
        
        if missing_fields:
            prompt = f"""
            Complete os seguintes campos para um documento {template_info['name']}:
            
            Dados fornecidos:
            {request.case_data}
            
            Campos a completar:
            {list(missing_fields)}
            
            Forneça sugestões apropriadas e profissionais para cada campo faltante,
            considerando o contexto jurídico e investigativo.
            
            Responda em formato JSON com os campos faltantes.
            """
            
            ai_result = await ai_orchestrator.analyze_with_provider(
                request.ai_provider,
                prompt,
                context=f"Geração de documento: {template_info['name']}"
            )
            
            if ai_result['success']:
                # Tentar parsear resposta da IA
                try:
                    import json
                    ai_suggestions = json.loads(ai_result['response'])
                    request.case_data.update(ai_suggestions)
                except:
                    # Se não for JSON válido, adicionar como texto
                    request.case_data['ai_suggestions'] = ai_result['response']
    
    # Gerar conteúdo do documento com IA
    prompt = f"""
    Gere o conteúdo completo para um documento: {template_info['name']}
    
    Descrição: {template_info['description']}
    
    Estrutura do documento:
    {template_info['structure']['secoes']}
    
    Dados do caso:
    {request.case_data}
    
    Gere um documento profissional, tecnicamente correto e juridicamente fundamentado.
    Use linguagem formal apropriada para contexto jurídico brasileiro.
    """
    
    content_result = await ai_orchestrator.analyze_with_provider(
        request.ai_provider,
        prompt,
        context="Geração de documento jurídico"
    )
    
    if not content_result['success']:
        raise HTTPException(status_code=500, detail="Erro ao gerar conteúdo do documento")
    
    # Criar documento Word
    doc = Document()
    
    # Título
    title = doc.add_heading(template_info['name'], level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Adicionar conteúdo gerado pela IA
    doc.add_paragraph(content_result['response'])
    
    # Adicionar rodapé com informações
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = f"Gerado em {datetime.now().strftime('%d/%m/%Y às %H:%M')} | AP Elite - Sistema ATHENA"
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Salvar documento
    filename = f"{request.template_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    filepath = os.path.join(GENERATED_DIR, filename)
    doc.save(filepath)
    
    # Salvar registro no banco
    doc_record = {
        "template_id": request.template_id,
        "template_name": template_info['name'],
        "case_data": request.case_data,
        "filename": filename,
        "filepath": filepath,
        "ai_provider": request.ai_provider,
        "ai_model": content_result['model'],
        "generated_at": datetime.now().isoformat()
    }
    
    result = await db.generated_documents.insert_one(doc_record)
    doc_record['id'] = str(result.inserted_id)
    
    return {
        "success": True,
        "document": doc_record,
        "download_url": f"/api/templates/download/{str(result.inserted_id)}"
    }

@router.get("/download/{document_id}")
async def download_document(document_id: str):
    """Download de documento gerado"""
    from fastapi.responses import FileResponse
    from bson import ObjectId
    
    try:
        doc = await db.generated_documents.find_one({"_id": ObjectId(document_id)})
    except:
        raise HTTPException(status_code=400, detail="ID inválido")
    
    if not doc:
        raise HTTPException(status_code=404, detail="Documento não encontrado")
    
    if not os.path.exists(doc['filepath']):
        raise HTTPException(status_code=404, detail="Arquivo não encontrado")
    
    return FileResponse(
        doc['filepath'],
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        filename=doc['filename']
    )

@router.get("/generated/list")
async def list_generated_documents(limit: int = 50):
    """Lista documentos gerados"""
    cursor = db.generated_documents.find().sort("generated_at", -1).limit(limit)
    documents = await cursor.to_list(length=limit)
    
    for doc in documents:
        doc['id'] = str(doc.pop('_id'))
    
    return {
        "documents": documents,
        "total": len(documents)
    }

@router.post("/ai-draft")
async def create_ai_draft(
    document_type: str,
    case_description: str,
    specific_requirements: Optional[str] = None
):
    """Cria rascunho de documento usando IA de forma livre"""
    
    prompt = f"""
    Crie um rascunho de documento jurídico/investigativo:
    
    Tipo: {document_type}
    Descrição do caso: {case_description}
    Requisitos específicos: {specific_requirements if specific_requirements else 'Nenhum'}
    
    Gere um documento completo, profissional e tecnicamente correto,
    seguindo as melhores práticas do direito brasileiro e perícia criminal.
    """
    
    # Usar consenso de múltiplos provedores para documentos importantes
    result = await ai_orchestrator.consensus_analysis(prompt)
    
    # Salvar rascunho
    draft_record = {
        "document_type": document_type,
        "case_description": case_description,
        "specific_requirements": specific_requirements,
        "draft_content": result['consensus'],
        "individual_analyses": result['individual_analyses'],
        "created_at": datetime.now().isoformat()
    }
    
    draft_result = await db.document_drafts.insert_one(draft_record)
    draft_record['id'] = str(draft_result.inserted_id)
    
    return {
        "success": True,
        "draft": draft_record
    }

@router.post("/improve")
async def improve_document(
    document_text: str,
    improvement_focus: str = 'general'
):
    """Melhora documento existente usando IA"""
    
    focus_prompts = {
        'general': 'melhore a clareza, coerência e profissionalismo',
        'legal': 'fortaleça os fundamentos jurídicos e citações legais',
        'technical': 'aprimore os aspectos técnicos e periciais',
        'grammar': 'corrija gramática, ortografia e formatação'
    }
    
    prompt = f"""
    Analise e melhore o seguinte documento:
    
    {document_text}
    
    Foco da melhoria: {focus_prompts.get(improvement_focus, focus_prompts['general'])}
    
    Forneça o documento melhorado mantendo a estrutura original,
    mas aprimorando qualidade, clareza e profissionalismo.
    """
    
    result = await ai_orchestrator.intelligent_analysis(
        'document_analysis',
        prompt
    )
    
    return {
        "success": result['success'],
        "original_length": len(document_text),
        "improved_document": result['response'],
        "provider": result['provider'],
        "model": result['model']
    }

@router.get("/statistics")
async def get_template_statistics():
    """Estatísticas de uso de templates"""
    
    total_generated = await db.generated_documents.count_documents({})
    total_drafts = await db.document_drafts.count_documents({})
    
    # Documentos por template
    pipeline = [
        {"$group": {
            "_id": "$template_id",
            "count": {"$sum": 1}
        }}
    ]
    
    by_template = []
    async for doc in db.generated_documents.aggregate(pipeline):
        template_info = DOCUMENT_TEMPLATES.get(doc['_id'], {})
        by_template.append({
            "template_id": doc['_id'],
            "template_name": template_info.get('name', 'Desconhecido'),
            "count": doc['count']
        })
    
    return {
        "total_templates": len(DOCUMENT_TEMPLATES),
        "total_generated": total_generated,
        "total_drafts": total_drafts,
        "by_template": by_template
    }
